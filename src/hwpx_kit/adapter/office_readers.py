"""PDF·DOCX·XLSX 텍스트 추출 — 순수 파이썬 (kordoc 흡수, 2026-07-10).

Node/kordoc 의존을 '한글 없는 환경의 구형 .hwp' 하나로 좁히기 위한 자체 리더.
출력은 read 명령의 content로 그대로 나가는 Markdown 유사 텍스트.
"""
from __future__ import annotations


def read_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        pages.append(f"## p.{i + 1}\n\n{text}" if len(reader.pages) > 1 else text)
    return "\n\n".join(pages)


def read_docx(path: str) -> str:
    import docx

    doc = docx.Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
    return "\n".join(parts)


def read_xlsx(path: str) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"## {ws.title}")
        for row in ws.iter_rows(values_only=True):
            if any(v is not None for v in row):
                parts.append("| " + " | ".join("" if v is None else str(v) for v in row) + " |")
    wb.close()
    return "\n".join(parts)
