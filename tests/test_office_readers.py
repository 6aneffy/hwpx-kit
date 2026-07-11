"""PDF·DOCX·XLSX 읽기 — kordoc 흡수 (순수 파이썬, Node 불필요).

kordoc은 'Mac 등 한글 없는 환경의 구형 .hwp' 전용 조건부 옵션으로 강등.
"""
import pytest

from hwpx_kit.commands.read import run_read

# 최소 유효 PDF (한 페이지, "Hello PDF" 텍스트)
_MINIMAL_PDF = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj
4 0 obj<</Length 44>>stream
BT /F1 24 Tf 72 720 Td (Hello PDF) Tj ET
endstream
endobj
5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
xref
0 6
0000000000 65535 f
trailer<</Size 6/Root 1 0 R>>
startxref
0
%%EOF
"""


@pytest.fixture(autouse=True)
def no_kordoc(monkeypatch):
    """kordoc을 차단해 순수 파이썬 경로임을 증명."""
    import hwpx_kit.adapter.kordoc_engine as ke

    monkeypatch.setattr(ke.shutil, "which", lambda name: None)


def test_read_docx(tmp_path):
    import docx

    d = docx.Document()
    d.add_paragraph("워드 본문 문장")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "라벨"
    t.rows[0].cells[1].text = "값"
    path = str(tmp_path / "doc.docx")
    d.save(path)

    result = run_read(path)
    assert "워드 본문 문장" in result["content"]
    assert "라벨" in result["content"] and "값" in result["content"]


def test_read_xlsx(tmp_path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "예산"
    ws["A1"] = "항목"
    ws["B1"] = 12340
    path = str(tmp_path / "data.xlsx")
    wb.save(path)

    result = run_read(path)
    assert "항목" in result["content"]
    assert "12340" in result["content"]
    assert "예산" in result["content"]  # 시트명


def test_read_pdf(tmp_path):
    path = str(tmp_path / "doc.pdf")
    with open(path, "wb") as fh:
        fh.write(_MINIMAL_PDF)
    result = run_read(path)
    assert "Hello PDF" in result["content"]


def test_read_hwp_without_kordoc_guides_convert(tmp_path):
    """kordoc 없는 환경에서 .hwp — convert 안내가 담긴 오류."""
    f = tmp_path / "old.hwp"
    f.write_bytes(b"x")
    with pytest.raises(RuntimeError, match="convert"):
        run_read(str(f))
